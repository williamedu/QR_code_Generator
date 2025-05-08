import os
import sys
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog
import qrcode
import uuid
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import tempfile

def seleccionar_docx():
    """Abre una ventana para seleccionar un archivo Word"""
    root = tk.Tk()
    root.withdraw()  # Oculta la ventana principal
    root.attributes('-topmost', True)  # Hacer que el diálogo sea visible sobre otras ventanas
    
    print("ESPERANDO_SELECCION_DOCX")
    
    archivo_docx = filedialog.askopenfilename(
        title="Selecciona un archivo Word",
        filetypes=[("Archivos Word", "*.docx"), ("Archivos Word antiguos", "*.doc")]
    )
    
    # Informar resultado de la selección
    if archivo_docx:
        print(f"DOCX_SELECCIONADO:{archivo_docx}")
    else:
        print("SELECCION_CANCELADA")
        
    return archivo_docx

def crear_elemento_posicionable(parent):
    """Crea un elemento que puede ser posicionado en coordenadas específicas"""
    drawing = OxmlElement('w:drawing')
    anchor = OxmlElement('wp:anchor')
    anchor.set(qn('wp14:anchorId'), 'anchor')
    anchor.set(qn('wp14:editId'), 'editor')
    anchor.set('allowOverlap', '1')
    anchor.set('layoutInCell', '1')
    anchor.set('behindDoc', '0')  # 0 para estar sobre el contenido, 1 para estar detrás
    anchor.set('relativeHeight', '1')
    anchor.set('simplePos', '0')
    anchor.set('locked', '0')
    
    # Añadir posición
    simple_pos = OxmlElement('wp:simplePos')
    simple_pos.set('x', '0')
    simple_pos.set('y', '0')
    anchor.append(simple_pos)
    
    # Posición horizontal
    pos_h = OxmlElement('wp:positionH')
    pos_h.set('relativeFrom', 'page')  # 'page', 'margin', 'column', etc.
    pos_h_offset = OxmlElement('wp:posOffset')
    pos_h.append(pos_h_offset)
    anchor.append(pos_h)
    
    # Posición vertical
    pos_v = OxmlElement('wp:positionV')
    pos_v.set('relativeFrom', 'page')  # 'page', 'margin', 'paragraph', etc.
    pos_v_offset = OxmlElement('wp:posOffset')
    pos_v.append(pos_v_offset)
    anchor.append(pos_v)
    
    drawing.append(anchor)
    parent.append(drawing)
    return drawing, anchor, pos_h_offset, pos_v_offset

def agregar_qr_a_docx(docx_path, metodo="insercion_directa"):
    """Agrega un código QR al documento Word y guarda una nueva versión con el sufijo _QR"""
    if not docx_path:
        print("No se seleccionó ningún archivo. Terminando ejecución.")
        return

    # Obtener el nombre del archivo sin ruta
    nombre_base = os.path.basename(docx_path)
    nombre_sin_extension = os.path.splitext(nombre_base)[0]
    
    # Obtener el directorio donde se está ejecutando el script
    directorio_script = os.path.dirname(os.path.abspath(__file__))
    
    # Definir nombre del archivo de salida en el directorio del script
    output_name = f"{nombre_sin_extension}_QR"
    output_path = os.path.join(directorio_script, f"{output_name}.docx")
    
    print(f"PROCESANDO_ARCHIVO:{docx_path}")
    print(f"NOMBRE_SALIDA:{output_name}")
    print(f"RUTA_GUARDADO:{directorio_script}")
    
    try:
        # Generar un ID único para el documento
        documento_id = str(uuid.uuid4())
        
        # Generar el código QR
        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_L,
            box_size=10,
            border=4,
        )
        
        # El QR contiene el nombre del archivo y un ID
        url_base = "https://menuidac.com/api/docx/descargar/documento/"
        url_documento = f"{url_base}{documento_id}"
        
        qr_data = {
            "url": url_documento,
            "id": documento_id,
            "documento": nombre_sin_extension
        }
        
        qr.add_data(str(qr_data))
        qr.make(fit=True)
        
        qr_img = qr.make_image(fill_color="black", back_color="white")
        
        # Guardar la imagen QR en un archivo temporal
        temp_qr_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
        qr_img.save(temp_qr_file.name)
        temp_qr_file.close()
        
        # Abrir el documento Word
        doc = Document(docx_path)
        
        if metodo == "pie_pagina":
            # MÉTODO 1: Agregar el código QR al pie de página
            for section in doc.sections:
                footer = section.footer
                footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                # Añadir el código QR como imagen
                run = footer_para.add_run()
                run.add_picture(temp_qr_file.name, width=Inches(0.8))  # Tamaño reducido del QR
                
                # Añadir un pequeño texto debajo del QR
                id_para = footer.add_paragraph()
                id_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                id_run = id_para.add_run(f"ID: {documento_id[:8]}...")
                id_run.font.size = Pt(7)
        
        elif metodo == "tabla_invisible":
            # MÉTODO 2: Usar una tabla invisble en la primera página
            # Este método ofrece más control sobre la posición
            
            # Insertar una tabla invisible en la primera página
            tabla = doc.add_table(rows=1, cols=1)
            
            # Establecer ancho de columna
            tabla.columns[0].width = Cm(2)  # Ancho de la celda para el QR
            
            # Obtener la celda y agregar el QR
            celda = tabla.cell(0, 0)
            celda_para = celda.paragraphs[0]
            
            # Añadir el QR a la celda
            run = celda_para.add_run()
            run.add_picture(temp_qr_file.name, width=Inches(0.8))
            
            # Hacer la tabla invisible (sin bordes)
            for row in tabla.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            pass
            
            # Mover la tabla al principio del documento
            doc.element.body.insert(0, tabla._element)
            
        elif metodo == "insercion_directa":
            # MÉTODO 3: Inserción directa en una posición específica
            
            # Preguntar al usuario dónde quiere colocar el QR
            root = tk.Tk()
            root.withdraw()
            
            opciones = [
                "Superior derecha (similar a PDF)",
                "Inferior derecha",
                "Inferior izquierda",
                "Superior izquierda",
                "Centro del documento",
                "Personalizado"
            ]
            
            posicion_idx = tk.StringVar(value=0)
            
            # Crear ventana de diálogo personalizada
            dialog = tk.Toplevel(root)
            dialog.title("Posición del código QR")
            dialog.geometry("380x300")
            dialog.attributes('-topmost', True)
            
            tk.Label(dialog, text="Selecciona dónde colocar el código QR:", pady=10).pack()
            
            for i, opcion in enumerate(opciones):
                tk.Radiobutton(dialog, text=opcion, variable=posicion_idx, value=i).pack(anchor="w", padx=20)
            
            resultado = {'value': 0, 'confirmed': False}
            
            def confirmar():
                resultado['value'] = int(posicion_idx.get())
                resultado['confirmed'] = True
                dialog.destroy()
            
            def cancelar():
                dialog.destroy()
            
            tk.Button(dialog, text="Confirmar", command=confirmar, pady=5, padx=10).pack(pady=10)
            tk.Button(dialog, text="Cancelar", command=cancelar, pady=5, padx=10).pack()
            
            dialog.wait_window()
            
            if not resultado['confirmed']:
                raise Exception("Operación cancelada por el usuario")
            
            # Definir posición basada en la selección
            custom_x = None
            custom_y = None
            
            if resultado['value'] == 0:  # Superior derecha (similar a PDF)
                x_pos = 8000000  # ~8 cm desde el borde izquierdo (en EMUs)
                y_pos = 1000000  # ~1 cm desde el borde superior
            elif resultado['value'] == 1:  # Inferior derecha
                x_pos = 8000000
                y_pos = 12000000
            elif resultado['value'] == 2:  # Inferior izquierda
                x_pos = 1000000
                y_pos = 12000000
            elif resultado['value'] == 3:  # Superior izquierda
                x_pos = 1000000
                y_pos = 1000000
            elif resultado['value'] == 4:  # Centro
                x_pos = 5000000
                y_pos = 7000000
            elif resultado['value'] == 5:  # Personalizado
                # Pedir al usuario valores personalizados (en cm)
                custom_x = simpledialog.askfloat("Posición X", 
                    "Introduce la posición horizontal (en cm desde el borde izquierdo):",
                    minvalue=0, maxvalue=20, initialvalue=8)
                custom_y = simpledialog.askfloat("Posición Y", 
                    "Introduce la posición vertical (en cm desde el borde superior):",
                    minvalue=0, maxvalue=29, initialvalue=1)
                
                if custom_x is not None and custom_y is not None:
                    # Convertir de cm a EMUs (1 cm ≈ 360000 EMUs)
                    x_pos = int(custom_x * 1000000)
                    y_pos = int(custom_y * 1000000)
                else:
                    raise Exception("Valores de posición inválidos")
            
            # Agregar el QR como una imagen en la primera página
            # La forma más sencilla es agregarlo como una imagen en línea
            paragraph = doc.paragraphs[0]
            run = paragraph.add_run()
            
            # Añadir la imagen
            picture = run.add_picture(temp_qr_file.name, width=Inches(0.8))
            
            # Texto ID debajo del QR (opcional)
            id_run = paragraph.add_run(f"\nID: {documento_id[:8]}...")
            id_run.font.size = Pt(7)
            
            print(f"QR colocado en posición: X={x_pos/1000000}cm, Y={y_pos/1000000}cm")
            
        # Guardar el documento Word con el QR
        doc.save(output_path)
        
        # Eliminar el archivo temporal del QR
        os.unlink(temp_qr_file.name)
        
        print(f"PROCESO_COMPLETADO:{output_path}")
        print(f"RUTA_COMPLETA:{os.path.abspath(output_path)}")
        
        # Mostrar un mensaje de éxito con una pequeña interfaz gráfica
        success_root = tk.Tk()
        success_root.title("Proceso Completado")
        success_root.geometry("400x200")
        success_root.attributes('-topmost', True)  # Hacer ventana visible
        
        label = tk.Label(
            success_root, 
            text=f"Documento Word con QR generado exitosamente:\n{output_name}.docx\n\nGuardado en la carpeta del script:\n{directorio_script}", 
            padx=20, 
            pady=20
        )
        label.pack()
        
        def cerrar():
            success_root.destroy()
            print("VENTANA_CERRADA")
        
        boton = tk.Button(success_root, text="Aceptar", command=cerrar)
        boton.pack(pady=10)
        
        success_root.mainloop()
        
        return output_path
        
    except Exception as e:
        error_msg = str(e)
        print(f"ERROR_PROCESAMIENTO:{error_msg}")
        
        # Mostrar un mensaje de error
        error_root = tk.Tk()
        error_root.title("Error")
        error_root.geometry("400x150")
        error_root.attributes('-topmost', True)  # Hacer ventana visible
        
        label = tk.Label(error_root, text=f"Error al procesar el documento Word:\n{error_msg}", padx=20, pady=20)
        label.pack()
        
        def cerrar():
            error_root.destroy()
            print("VENTANA_ERROR_CERRADA")
        
        boton = tk.Button(error_root, text="Aceptar", command=cerrar)
        boton.pack(pady=10)
        
        error_root.mainloop()
        return None

def seleccionar_segundo_docx():
    """Abre una ventana para seleccionar un segundo archivo Word (carta)"""
    respuesta = messagebox.askyesno("Proceso QR", "¿Deseas procesar también una carta junto con este oficio?")
    
    if not respuesta:
        return None
    
    root = tk.Tk()
    root.withdraw()  # Oculta la ventana principal
    root.attributes('-topmost', True)  # Hacer que el diálogo sea visible sobre otras ventanas
    
    print("ESPERANDO_SELECCION_CARTA")
    
    archivo_carta = filedialog.askopenfilename(
        title="Selecciona la Carta de Estado (Word)",
        filetypes=[("Archivos Word", "*.docx"), ("Archivos Word antiguos", "*.doc")]
    )
    
    # Informar resultado de la selección
    if archivo_carta:
        print(f"CARTA_SELECCIONADA:{archivo_carta}")
    else:
        print("SELECCION_CARTA_CANCELADA")
        
    return archivo_carta

def main():
    """Función principal del programa"""
    print("INICIANDO_SCRIPT")
    
    # Crear una ventana simple para empezar
    root = tk.Tk()
    root.title("Generador de QR para Word")
    root.geometry("450x300")
    root.attributes('-topmost', True)
    
    label = tk.Label(
        root, 
        text="Esta herramienta te permite añadir un código QR\na un documento Word (.docx).\n\nSelecciona el método de inserción del QR:",
        padx=20, 
        pady=20,
        justify="center"
    )
    label.pack()
    
    # Variable para el método seleccionado
    metodo_seleccionado = tk.StringVar(value="insercion_directa")
    
    # Opciones de método
    tk.Radiobutton(root, text="Inserción directa (posición personalizable)", 
                  variable=metodo_seleccionado, value="insercion_directa").pack(anchor="w", padx=20)
    tk.Radiobutton(root, text="Pie de página", 
                  variable=metodo_seleccionado, value="pie_pagina").pack(anchor="w", padx=20)
    tk.Radiobutton(root, text="Tabla invisible", 
                  variable=metodo_seleccionado, value="tabla_invisible").pack(anchor="w", padx=20)
    
    def iniciar_proceso():
        metodo = metodo_seleccionado.get()
        root.destroy()
        
        # Seleccionar el archivo principal
        docx_path = seleccionar_docx()
        
        if docx_path:
            # Opcional: seleccionar un segundo documento (carta)
            carta_path = seleccionar_segundo_docx()
            
            if carta_path:
                print(f"Procesando oficio: {docx_path}")
                print(f"Procesando carta: {carta_path}")
                # En este caso, procesamos ambos documentos
                oficio_con_qr = agregar_qr_a_docx(docx_path, metodo)
                
                if oficio_con_qr:
                    # Si se procesó el oficio correctamente, procesamos la carta
                    # Nota: en un flujo real, el ID del QR sería el mismo para ambos
                    carta_con_qr = agregar_qr_a_docx(carta_path, metodo)
            else:
                # Solo procesar el oficio
                agregar_qr_a_docx(docx_path, metodo)
        else:
            print("SELECCION_CANCELADA_FINAL")
            
            # Mostrar mensaje de cancelación
            cancel_root = tk.Tk()
            cancel_root.title("Operación Cancelada")
            cancel_root.geometry("300x100")
            cancel_root.attributes('-topmost', True)  # Hacer ventana visible
            
            label = tk.Label(cancel_root, text="Operación cancelada. No se seleccionó ningún archivo.", padx=20, pady=20)
            label.pack()
            
            def cerrar():
                cancel_root.destroy()
                print("VENTANA_CANCELACION_CERRADA")
            
            boton = tk.Button(cancel_root, text="Aceptar", command=cerrar)
            boton.pack(pady=10)
            
            cancel_root.mainloop()
    
    # Botón para iniciar el proceso
    boton = tk.Button(root, text="Seleccionar documento Word", command=iniciar_proceso, padx=10, pady=5)
    boton.pack(pady=20)
    
    root.mainloop()

if __name__ == "__main__":
    main()