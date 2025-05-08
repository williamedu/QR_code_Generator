import os
import sys
import uuid
import datetime
import json
import tempfile
import shutil
from io import BytesIO
import pymysql
import qrcode
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pdf2docx import Converter
import pythoncom  # Para COM (utilizado con win32com si es necesario)

# Importaciones para convertir DOCX a PDF (para visualización consistente)
import subprocess
from docx2pdf import convert

# Importaciones para la API REST
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from werkzeug.utils import secure_filename

# 📌 Configuración de la base de datos
DB_HOST = "americastowersimulator.c14c80caytj6.us-east-1.rds.amazonaws.com"
DB_USER = "admin"
DB_PASSWORD = "Controlador2929"
DB_NAME = "simulador(unity-access)"

# Configuración de la aplicación
app = Flask(__name__)
CORS(app)  # Habilitar CORS para todas las rutas

app.config['UPLOAD_FOLDER'] = tempfile.gettempdir()  # Carpeta temporal para archivos subidos
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # Limitar tamaño de archivos a 16MB
app.config['OUTPUT_FOLDER'] = r"C:\Users\Administrator\Desktop\QR_code_Generator\documentos_word_procesados"  # Ruta específica para documentos Word

# Crear carpeta de salida si no existe
if not os.path.exists(app.config['OUTPUT_FOLDER']):
    os.makedirs(app.config['OUTPUT_FOLDER'])

# 📌 Función para conectar con la base de datos
def conectar_bd():
    try:
        conexion = pymysql.connect(
            host=DB_HOST,
            user=DB_USER,
            password=DB_PASSWORD,
            database=DB_NAME,
            cursorclass=pymysql.cursors.DictCursor
        )
        return conexion
    except pymysql.MySQLError as e:
        app.logger.error(f"Error al conectar con la base de datos: {e}")
        return None

def guardar_datos_bd(datos):
    """Guarda los datos en la base de datos"""
    conexion = conectar_bd()
    if not conexion:
        app.logger.error("No se pudo establecer conexión con la base de datos")
        return False
    
    try:
        with conexion.cursor() as cursor:
            # Insertar los datos en la tabla
            # Usamos la misma tabla que para PDFs pero con campo adicional que indica que es DOCX
            cursor.execute("""
            INSERT INTO documentos_qr (
                id, nombre_original, nombre_con_qr, s3_bucket, s3_key, 
                s3_url, tamano_archivo, qr_data, descripcion, metadata, tipo_archivo
            ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            """, (
                datos["id"],
                datos["nombre_original"],
                datos["nombre_con_qr"],
                "docs-qr-bucket",  # Nombre del bucket (simulado por ahora)
                datos["s3_key"],
                datos["s3_url"],
                datos["tamano_bytes"],
                json.dumps(datos["qr_data"]),
                datos["descripcion"],
                json.dumps(datos["metadata"]),
                "DOCX"  # Indicamos que es un documento Word
            ))
            
            conexion.commit()
            app.logger.info(f"Registro con ID {datos['id']} guardado exitosamente")
            return True
    except pymysql.MySQLError as e:
        app.logger.error(f"Error al insertar en la base de datos: {e}")
        return False
    finally:
        conexion.close()

def extraer_datos_carta_estado_docx(archivo_docx):
    """Extrae información relevante de la carta de estado en formato DOCX"""
    # Generar un UUID para el documento
    documento_id = str(uuid.uuid4())
    
    # Obtener el nombre del archivo
    nombre_original = secure_filename(archivo_docx.filename)
    
    # Leer información básica del DOCX
    docx_data = archivo_docx.read()
    archivo_docx.seek(0)  # Resetear el puntero para usos futuros
    
    # Crear un objeto BytesIO para leer el DOCX
    docx_bytes = BytesIO(docx_data)
    
    # Leer el documento con python-docx para extraer más información
    try:
        doc = Document(docx_bytes)
        # Podríamos extraer texto o metadatos específicos aquí si es necesario
        
        # Contar párrafos como aproximación del contenido
        num_paragrafos = len(doc.paragraphs)
        
        # Datos extraídos
        datos = {
            "id": documento_id,
            "nombre_original": nombre_original,
            "fecha_creacion": datetime.datetime.now().isoformat(),
            "tipo_documento": "Carta de Estado DOCX",
            "tamano_bytes": len(docx_data),
            "parrafos": num_paragrafos,
            "docx_data": docx_data  # Guardar los datos binarios para uso posterior
        }
        
        app.logger.info(f"Datos extraídos de DOCX con ID generado {documento_id}")
        return datos
    except Exception as e:
        app.logger.error(f"Error al leer el documento DOCX: {e}")
        raise

def generar_qr_con_datos(datos):
    """Genera un código QR con la información de la carta de estado"""
    # Crear la URL que iría en el QR con la IP y puerto del servidor
    url_base = "https://menuidac.com/api/docx/descargar/documento/"
    url_documento = f"{url_base}{datos['id']}"
    
    # Crear datos para el QR
    qr_data = {
        "url": url_documento,
        "id": datos["id"],
        "tipo": datos["tipo_documento"],
        "nombre": datos["nombre_original"],
        "fecha": datos["fecha_creacion"]
    }
    
    # Convertir a JSON
    qr_json = json.dumps(qr_data)
    
    # Generar el código QR
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_M,
        box_size=10,
        border=4,
    )
    
    # Añadir los datos al QR
    qr.add_data(qr_json)
    qr.make(fit=True)
    
    qr_img = qr.make_image(fill_color="black", back_color="white")
    
    # Guardar la imagen QR en un archivo temporal
    temp_qr_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
    qr_img.save(temp_qr_file.name)
    temp_qr_file.close()
    
    app.logger.info(f"QR generado con información del documento DOCX {datos['id']}")
    return temp_qr_file.name, qr_data

def agregar_qr_a_oficio_docx(oficio_data, qr_image_path, qr_data):
    """Agrega un código QR al oficio DOCX y retorna el nuevo documento"""
    # Nombre para el archivo resultante
    nombre_original = oficio_data["nombre_original"]
    nombre_sin_extension = os.path.splitext(nombre_original)[0]
    nombre_salida = f"{qr_data['id']}_{nombre_sin_extension}_con_QR.docx"
    ruta_salida = os.path.join(app.config['OUTPUT_FOLDER'], nombre_salida)
    
    try:
        # Crear un objeto BytesIO para leer el DOCX
        docx_bytes = BytesIO(oficio_data["docx_data"])
        
        # Abrir el documento con python-docx
        documento = Document(docx_bytes)
        
        # Obtener el último párrafo de la primera sección
        section = documento.sections[0]
        
        # Añadir un header o footer para el QR (elegimos footer)
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Añadir el código QR como imagen
        run = footer_para.add_run()
        run.add_picture(qr_image_path, width=Inches(0.8))  # Tamaño reducido del QR
        
        # Añadir un pequeño texto debajo del QR
        id_para = footer.add_paragraph()
        id_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        id_run = id_para.add_run(f"ID: {qr_data['id'][:8]}...")
        id_run.font.size = Pt(7)
        
        # Guardar el documento resultante
        documento.save(ruta_salida)
        
        # Eliminar el archivo temporal del QR
        os.unlink(qr_image_path)
        
        app.logger.info(f"Oficio DOCX con QR guardado en {ruta_salida}")
        return ruta_salida
        
    except Exception as e:
        app.logger.error(f"Error al agregar QR al DOCX: {str(e)}")
        # Limpiar
        if os.path.exists(qr_image_path):
            os.unlink(qr_image_path)
        return None

def procesar_archivos_docx(carta_data, oficio_data):
    """Procesa los dos archivos DOCX recibidos"""
    app.logger.info(f"Procesando carta DOCX: {carta_data['nombre_original']}")
    app.logger.info(f"Procesando oficio DOCX: {oficio_data['nombre_original']}")
    
    try:
        # 1. Generar QR con la información relevante
        qr_path, qr_data = generar_qr_con_datos(carta_data)
        
        # 2. Guardar una copia local de la carta de estado (incluir ID en el nombre)
        nombre_con_id = f"{carta_data['id']}_{carta_data['nombre_original']}"
        ruta_copia_carta = os.path.join(app.config['OUTPUT_FOLDER'], nombre_con_id)
        
        with open(ruta_copia_carta, 'wb') as f:
            f.write(carta_data['docx_data'])
        app.logger.info(f"Carta DOCX guardada en: {ruta_copia_carta}")
        
        # 3. Añadir el QR al oficio
        ruta_oficio_con_qr = agregar_qr_a_oficio_docx(oficio_data, qr_path, qr_data)
        
        if not ruta_oficio_con_qr:
            return {
                "error": "Error al agregar QR al oficio DOCX",
                "success": False
            }
        
        # 4. Preparar datos para la base de datos
        nombre_sin_extension = os.path.splitext(oficio_data['nombre_original'])[0]
        
        datos_bd = {
            "id": carta_data["id"],
            "nombre_original": nombre_con_id,  # Guardar con el nombre que incluye ID
            "nombre_con_qr": f"{carta_data['id']}_{nombre_sin_extension}_con_QR.docx",
            "s3_key": f"cartas_docx/{carta_data['id']}/{nombre_con_id}",
            "s3_url": f"https://menuidac.com/api/docx/descargar/documento/{carta_data['id']}",
            "tamano_bytes": carta_data["tamano_bytes"],
            "qr_data": qr_data,
            "descripcion": f"Carta de estado DOCX para oficio: {oficio_data['nombre_original']}",
            "metadata": {
                "oficio_relacionado": oficio_data['nombre_original'],
                "ruta_oficio_con_qr": ruta_oficio_con_qr,
                "fecha_procesamiento": datetime.datetime.now().isoformat()
            }
        }
        
        # 5. Guardar en base de datos
        bd_resultado = guardar_datos_bd(datos_bd)
        
        if bd_resultado:
            app.logger.info("Ambos archivos DOCX procesados y datos guardados correctamente")
            return {
                "carta_guardada": ruta_copia_carta,
                "oficio_con_qr": ruta_oficio_con_qr,
                "id_documento": carta_data["id"],
                "success": True
            }
        else:
            app.logger.error("No se pudieron guardar los datos en la base de datos")
            return {
                "error": "Error al guardar en la base de datos",
                "success": False
            }
            
    except Exception as e:
        app.logger.error(f"Error en el procesamiento de DOCX: {str(e)}")
        return {
            "error": str(e),
            "success": False
        }

# Conversión opcional entre DOCX y PDF para visualización
def convertir_docx_a_pdf(ruta_docx):
    """Convierte un archivo DOCX a PDF para visualización"""
    try:
        ruta_pdf = os.path.splitext(ruta_docx)[0] + ".pdf"
        convert(ruta_docx, ruta_pdf)
        return ruta_pdf
    except Exception as e:
        app.logger.error(f"Error al convertir DOCX a PDF: {str(e)}")
        return None

# Ruta de verificación de salud
@app.route('/api/docx/health', methods=['GET'])
def health_check():
    """Endpoint para verificar que la API está funcionando"""
    return jsonify({"status": "ok", "message": "API de DOCX funcionando correctamente"})

# Endpoint para procesar los DOCX
@app.route('/api/docx/procesar', methods=['POST'])
def procesar_docxs():
    """Endpoint para recibir y procesar dos archivos DOCX"""
    try:
        # Verificar que se hayan enviado los dos archivos
        if 'carta' not in request.files or 'oficio' not in request.files:
            return jsonify({"error": "Se requieren ambos archivos: 'carta' y 'oficio'", "success": False}), 400
        
        archivo_carta = request.files['carta']
        archivo_oficio = request.files['oficio']
        
        # Verificar que los archivos tengan nombres
        if archivo_carta.filename == '' or archivo_oficio.filename == '':
            return jsonify({"error": "Ambos archivos deben tener un nombre", "success": False}), 400
        
        # Verificar que los archivos sean DOCX
        if not archivo_carta.filename.lower().endswith(('.docx', '.doc')) or not archivo_oficio.filename.lower().endswith(('.docx', '.doc')):
            return jsonify({"error": "Ambos archivos deben ser documentos de Word (.docx o .doc)", "success": False}), 400
        
        # Verificar si son DOC (antiguo) y convertirlos a DOCX si es necesario
        # Este paso podría necesitar implementación adicional con win32com si se necesita manejar .DOC
        
        # Extraer datos de la carta de estado
        carta_data = extraer_datos_carta_estado_docx(archivo_carta)
        
        # Extraer datos del oficio
        oficio_data = {
            "nombre_original": secure_filename(archivo_oficio.filename),
            "docx_data": archivo_oficio.read()
        }
        
        # Procesar los archivos
        resultado = procesar_archivos_docx(carta_data, oficio_data)
        
        if resultado.get("success", False):
            # Ofrecer la descarga del oficio con QR
            oficio_con_qr = resultado.get("oficio_con_qr")
            
            # Reemplazar rutas locales con URLs relativas para la API
            nombre_oficio_qr = os.path.basename(oficio_con_qr)
            resultado["oficio_con_qr_url"] = f"/api/docx/descargar/{nombre_oficio_qr}"
            
            return jsonify(resultado), 200
        else:
            return jsonify(resultado), 500
            
    except Exception as e:
        app.logger.error(f"Error en el endpoint de DOCX: {str(e)}")
        return jsonify({"error": str(e), "success": False}), 500

# Endpoint para descargar archivos procesados por nombre
@app.route('/api/docx/descargar/<nombre_archivo>', methods=['GET'])
def descargar_archivo(nombre_archivo):
    """Permite descargar un archivo DOCX procesado por su nombre"""
    try:
        ruta_archivo = os.path.join(app.config['OUTPUT_FOLDER'], secure_filename(nombre_archivo))
        
        if not os.path.exists(ruta_archivo):
            app.logger.error(f"Archivo DOCX no encontrado: {ruta_archivo}")
            return jsonify({"error": "Archivo no encontrado", "success": False}), 404
            
        # Opción para convertir a PDF para visualización web (parámetro opcional)
        format_pdf = request.args.get('pdf', 'false').lower() == 'true'
        
        if format_pdf and ruta_archivo.lower().endswith('.docx'):
            pdf_path = convertir_docx_a_pdf(ruta_archivo)
            if pdf_path and os.path.exists(pdf_path):
                return send_file(pdf_path, as_attachment=True)
        
        return send_file(ruta_archivo, as_attachment=True)
        
    except Exception as e:
        app.logger.error(f"Error al descargar archivo DOCX: {str(e)}")
        return jsonify({"error": str(e), "success": False}), 500

# Endpoint para descargar por ID del documento
@app.route('/api/docx/descargar/documento/<documento_id>', methods=['GET'])
def descargar_por_id(documento_id):
    """Permite descargar un documento DOCX por su ID"""
    try:
        app.logger.info(f"Buscando documento DOCX con ID: {documento_id}")
        
        # Buscar en el directorio configurado cualquier archivo que contenga el ID
        directorio = app.config['OUTPUT_FOLDER']
        
        try:
            # Listar todos los archivos en el directorio
            archivos = os.listdir(directorio)
            app.logger.info(f"Total de archivos en el directorio: {len(archivos)}")
            
            # Filtrar archivos que contengan el ID y sean DOCX
            archivos_coincidentes = [archivo for archivo in archivos 
                                    if documento_id in archivo and 
                                    (archivo.lower().endswith('.docx') or archivo.lower().endswith('.doc'))]
            
            app.logger.info(f"Archivos DOCX que coinciden con el ID {documento_id}: {archivos_coincidentes}")
            
            if not archivos_coincidentes:
                app.logger.error(f"No se encontró ningún archivo DOCX con el ID: {documento_id}")
                return jsonify({"error": "Archivo no encontrado", "success": False}), 404
            
            # Usar el primer archivo que coincida
            archivo_encontrado = archivos_coincidentes[0]
            ruta_completa = os.path.join(directorio, archivo_encontrado)
            
            app.logger.info(f"Archivo DOCX encontrado: {ruta_completa}")
            
            # Verificar si el archivo existe
            if not os.path.exists(ruta_completa):
                app.logger.error(f"El archivo DOCX coincidente no existe: {ruta_completa}")
                return jsonify({"error": "Archivo existe en lista pero no en sistema", "success": False}), 404
            
            # Opción para convertir a PDF para visualización web (parámetro opcional)
            format_pdf = request.args.get('pdf', 'false').lower() == 'true'
            
            if format_pdf and ruta_completa.lower().endswith('.docx'):
                pdf_path = convertir_docx_a_pdf(ruta_completa)
                if pdf_path and os.path.exists(pdf_path):
                    return send_file(pdf_path, as_attachment=True)
            
            # Devolver el archivo
            return send_file(ruta_completa, as_attachment=True)
            
        except Exception as e:
            app.logger.error(f"Error al buscar en el directorio para DOCX: {str(e)}")
            return jsonify({"error": f"Error al buscar en el directorio: {str(e)}", "success": False}), 500
        
    except Exception as e:
        app.logger.error(f"Error al buscar documento DOCX por ID: {str(e)}")
        return jsonify({"error": str(e), "success": False}), 500

if __name__ == '__main__':
    # Configurar logging
    import logging
    logging.basicConfig(level=logging.INFO)
    
    # Puerto donde se ejecutará la API (distinto al de la API de PDF)
    puerto = 5001
    
    print(f"Iniciando API para documentos Word en http://0.0.0.0:{puerto}")
    print("Endpoints disponibles:")
    print(f"  - GET https://menuidac.com/api/docx/health - Verificar si la API está funcionando")
    print(f"  - POST https://menuidac.com/api/docx/procesar - Procesar los archivos DOCX")
    print(f"  - GET https://menuidac.com/api/docx/descargar/<nombre_archivo> - Descargar por nombre")
    print(f"  - GET https://menuidac.com/api/docx/descargar/documento/<documento_id> - Descargar por ID")
    
    # Usar waitress para producción
    from waitress import serve
    serve(app, host='0.0.0.0', port=puerto, threads=4)
